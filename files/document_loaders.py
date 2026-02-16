"""
Document Loaders for DDR Pipeline
==================================
Utilities to load and extract text from various document formats
"""

import os
from typing import Optional
import io


def load_pdf(file_path: str) -> str:
    """
    Load text from PDF file
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        Extracted text content
    """
    try:
        import PyPDF2
        
        text_content = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            print(f"Loading PDF: {os.path.basename(file_path)} ({num_pages} pages)")
            
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    text_content.append(text)
        
        full_text = "\n".join(text_content)
        print(f"✓ Extracted {len(full_text)} characters from PDF")
        return full_text
        
    except ImportError:
        raise ImportError("PyPDF2 not installed. Install with: pip install PyPDF2")
    except Exception as e:
        raise Exception(f"Error loading PDF: {str(e)}")


def load_docx(file_path: str) -> str:
    """
    Load text from DOCX file
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Extracted text content
    """
    try:
        from docx import Document
        
        print(f"Loading DOCX: {os.path.basename(file_path)}")
        
        doc = Document(file_path)
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
        
        full_text = "\n".join(text_content)
        print(f"✓ Extracted {len(full_text)} characters from DOCX")
        return full_text
        
    except ImportError:
        raise ImportError("python-docx not installed. Install with: pip install python-docx")
    except Exception as e:
        raise Exception(f"Error loading DOCX: {str(e)}")


def load_text(file_path: str) -> str:
    """
    Load text from TXT file
    
    Args:
        file_path: Path to TXT file
        
    Returns:
        Text content
    """
    try:
        print(f"Loading TXT: {os.path.basename(file_path)}")
        
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        print(f"✓ Loaded {len(text)} characters from TXT")
        return text
        
    except Exception as e:
        raise Exception(f"Error loading TXT: {str(e)}")


def load_document(file_path: str) -> str:
    """
    Auto-detect file type and load content
    
    Args:
        file_path: Path to document file
        
    Returns:
        Extracted text content
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    ext = os.path.splitext(file_path)[1].lower()
    
    loaders = {
        '.pdf': load_pdf,
        '.docx': load_docx,
        '.txt': load_text,
    }
    
    loader = loaders.get(ext)
    if loader is None:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {list(loaders.keys())}")
    
    return loader(file_path)


def save_text_output(text: str, output_path: str) -> None:
    """
    Save text content to file
    
    Args:
        text: Text content to save
        output_path: Path to output file
    """
    with open(output_path, 'w', encoding='utf-8') as file:
        file.write(text)
    print(f"✓ Saved output to: {output_path}")
